def summary(filename):
  from docx import Document
  from docx.shared import Pt
  import csv

  #Creates 'file' as a shortcut to opening the desired CSV document
  file = open(filename)

  rest_start_times = []
  rest_end_times = []
  rest_durations = []
  sleep_start_times = []
  sleep_end_times = []
  sleep_durations = []
  epoch_detected = False

  #Changes time from 12 hour time to 24 hour time (now incorperated in Brandon's code aswell)
  def time_24hr(time):
      time = time.lower()
      if 'pm' in time:
          time = time.strip(' pm')
          time = time.split(':')
          if int(time[0]) < 12:
              time[0] = int(time[0]) + 12
              time[0] = str(time[0])
          time = ":".join(time)
      elif 'am' in time:
          time = time.strip(' am')
          time = time.split(':')
          if int(time[0]) == 12:
              time[0] = 0
              time[0] = str(time[0])
          time = ":".join(time)
      return(time)

  def time_min(list):
      min_time = list[0]
      if type(min_time) == type('a'):
          min_time = [int(part) for part in min_time.split(":")]
      for value in range(len(list)):
          integer = list[value]
          if type(integer) == type('a'):
              integer = [int(part) for part in integer.split(":")]
          if integer[0] < min_time[0] and integer[0] > 12:
              min_time = integer
          elif integer[0] == min_time[0]:
              if integer[1] < min_time[1]:
                  min_time = integer
      return(min_time)

  def time_max(list):
      max_time = list[0]
      if type(max_time) == type('a'):
          max_time = [int(part) for part in max_time.split(":")]
      for value in range(len(list)):
          integer = list[value]
          if type(integer) == type('a'):
              integer = [int(part) for part in integer.split(":")]
          if integer[0] > max_time[0]:
              max_time = integer
          elif integer[0] == max_time[0]:
              if integer[1] > max_time[1]:
                  max_time = integer
      return(max_time)

  def to_secs(my_list):
      times_in_secs = []
      for i in range(len(my_list)):
          if type(my_list[i]) == type('a'):
              my_list[i] = [int(part) for part in my_list[i].split(":")]
          if my_list[i][0] < 12:
              my_list[i][0] = my_list[i][0] + 24
          times_in_secs.append((my_list[i][0] * 60 * 60) + my_list[i][1] * 60)
      return(times_in_secs)

  def from_secs(value):
      time = [0,0,0]
      time[2] = int(value%60)
      value = (value - value%60)/60
      time[1] = int(value%60)
      value = (value - value%60)/60
      time[0] = int(value)
      return(time)

  def time_avg(list):
      times = list.copy()
      times_in_secs = to_secs(list)
      sum_of_times = 0
      for counter in range(len(times)):
          sum_of_times = sum_of_times + times_in_secs[counter]
      avg_time = sum_of_times/len(times_in_secs)
      return(from_secs(avg_time))

  def add_zero(string):
      split_string = [str(part) for part in string.split(":")]
      for counter in range(len(split_string)):
          if len(split_string[counter]) < 2:
              a = split_string[counter]
              split_string[counter] = '0'+ a
      string = ":".join(split_string)
      return(string)


  for line in file:
      if "Epoch-by-Epoch Data" in line:
          epoch_detected = True
      line = line.strip('"')
      line = line.strip('",\n')
      line = line.split('","')

  #extracts row by row of time in bed/rest times
      if 'REST' in line and epoch_detected == False:
          time1 = time_24hr(line[4])
          rest_start_times.append(time1)

          time2 = time_24hr(line[7])
          rest_end_times.append(time2)

  #Calculate duration of hours in bed
          time1 = [int(part) for part in time1.split(":")]
          time2 = [int(part) for part in time2.split(":")]
          if time1 != [0.0]:
              time3 = [0,0,0]
              if time1[0] > time2[0]:
                  hour_difference = 24 - time1[0]
                  time3[0] = time3[0] + hour_difference
              time3[0] = time3[0] + time2[0]
              minute_difference = time1[1] - time2[1]
              if minute_difference < 0:
                  time3[0] = time3[0] - 1
                  time3[1] = 60 + minute_difference
              else:
                  time3[1] = minute_difference
              rest_durations.append(time3)

  #Extracting row by row of recorded sleep times
      if 'SLEEP' in line and epoch_detected == False:
          stime1 = time_24hr(line[4])
          sleep_start_times.append(stime1)

  #Append end time
          stime2 = time_24hr(line[7])
          sleep_end_times.append(stime2)

  #Calculate duration and append to sleep_durations
          stime1 = [int(part) for part in stime1.split(":")]
          stime2 = [int(part) for part in stime2.split(":")]
          stime3 = [0,0,0]
          if stime1 != [0.0]:
              if stime1[0] > stime2[0]:
                  hour_difference = 24 - stime1[0]
                  stime3[0] = stime3[0] + hour_difference
              stime3[0] = stime3[0] + stime2[0]
              minute_difference = stime1[1] - stime2[1]
              if minute_difference < 0:
                  stime3[0] = stime3[0] - 1
                  stime3[1] = 60 + minute_difference
              else:
                  stime3[1] = minute_difference
              sleep_durations.append(stime3)

  #Minimums for all except time
      elif 'Sleep Summary' in line and 'Minimum(n)' in line:
          min_onset_latency = line[20]
          min_efficiency = line[21]
          min_waso = line[23]
          min_awakenings = line[25]

  #Maximums for all except time
      elif 'Sleep Summary' in line and 'Maximum(n)' in line:
          max_onset_latency = line[20]
          max_efficiency = line[21]
          max_waso = line[23]
          max_awakenings = line[25]

  #Averages for all except time
      elif 'Sleep Summary' in line and 'Average(n)' in line:
          avg_onset_latency = line[20]
          avg_efficiency = line[21]
          avg_waso = line[23]
          avg_awakenings = line[25]

  #Minimums for time values
  min_bedtime = add_zero(":".join([str(part) for part in time_min(rest_start_times)]))
  min_get_up_time = add_zero(":".join([str(part) for part in time_min(rest_end_times)]))
  min_time_in_bed = add_zero(":".join([str(part) for part in time_min(rest_durations)]))
  min_total_sleep_hours = add_zero(":".join([str(part) for part in time_min(sleep_durations)]))

  #Maximums for time values
  max_bedtime = add_zero(":".join([str(part) for part in time_max(rest_start_times)]))
  max_get_up_time = add_zero(":".join([str(part) for part in time_max(rest_end_times)]))
  max_time_in_bed = add_zero(":".join([str(part) for part in time_max(rest_durations)]))
  max_total_sleep_hours = add_zero(":".join([str(part) for part in time_max(sleep_durations)]))

  #Averges for time values
  avg_bedtime = add_zero(":".join([str(part) for part in time_avg(rest_start_times)]))
  avg_get_up_time = add_zero(":".join([str(part) for part in time_avg(rest_end_times)]))
  avg_time_in_bed = add_zero(":".join([str(part) for part in time_avg(rest_durations)]))
  avg_total_sleep_hours = add_zero(":".join([str(part) for part in time_avg(sleep_durations)]))

  #Export to the summary document as a csv
  out = open("summary.csv","w")
  out.write("SUMMARY STATISTICS"+'\n')
  out.write('\n')
  out.write(''+","'Bed Time'+","'Get Up Time'+","'Time in Bed (hours)'+","'Total Sleep Time (hours)'+","'Onset Latency (minutes)'+","'Sleep Efficiency (Percent)'+","'WASO (minutes)'+","'#Awak'+'\n')
  out.write('Min'+",")
  out.write(min_bedtime+","+min_get_up_time+","+min_time_in_bed+","+min_total_sleep_hours+","+min_onset_latency+","+min_efficiency+","+min_waso+","+min_awakenings+'\n')
  out.write('Max'+",")
  out.write(max_bedtime+","+max_get_up_time+","+max_time_in_bed+","+max_total_sleep_hours+","+max_onset_latency+","+max_efficiency+","+max_waso+","+max_awakenings+'\n')
  out.write('Avg'+",")
  out.write(avg_bedtime+","+avg_get_up_time+","+avg_time_in_bed+","+avg_total_sleep_hours+","+avg_onset_latency+","+avg_efficiency+","+avg_waso+","+avg_awakenings)
  out.close()


  #Exports as a word document
  #Word document with a title
  document = Document()
  document.add_heading('Sleep Summary Statistics', 0)

  #blank line
  blank_line = document.add_paragraph('')

  #Creating a table
  table = document.add_table(rows=4, cols=9)
  table.style = 'Table Grid'

  #Column titles
  title_cell = table.cell(0, 1)
  title_cell.text = 'Bed Time'
  blank_line = title_cell.add_paragraph('')
  title_cell = table.cell(0, 2)
  title_cell.text = 'Get Up Time'
  blank_line = title_cell.add_paragraph('')
  title_cell = table.cell(0, 3)
  title_cell.text = 'Time in Bed (hours)'
  blank_line = title_cell.add_paragraph('')
  title_cell = table.cell(0, 4)
  title_cell.text = 'Total Sleep Time (hours)'
  blank_line = title_cell.add_paragraph('')
  title_cell = table.cell(0, 5)
  title_cell.text = 'Onset Latency (minutes)'
  blank_line = title_cell.add_paragraph('')
  title_cell = table.cell(0, 6)
  title_cell.text = 'Sleep Efficiency (Percent)'
  blank_line = title_cell.add_paragraph('')
  title_cell = table.cell(0, 7)
  title_cell.text = 'WASO (minutes)'
  blank_line = title_cell.add_paragraph('')
  title_cell = table.cell(0, 8)
  title_cell.text = '#Awak.'
  blank_line = title_cell.add_paragraph('')

  #Row titles
  title_cell = table.cell(1, 0)
  title_cell.text = 'Min'
  title_cell = table.cell(2, 0)
  title_cell.text = 'Max'
  title_cell = table.cell(3, 0)
  title_cell.text = 'Avg'

  #Row 1 data
  cell = table.cell(1, 1)
  cell.text = min_bedtime
  blank_line = cell.add_paragraph('')
  cell = table.cell(1, 2)
  cell.text = min_get_up_time
  cell = table.cell(1, 3)
  cell.text = min_time_in_bed
  cell = table.cell(1, 4)
  cell.text = min_total_sleep_hours
  cell = table.cell(1, 5)
  cell.text = min_onset_latency
  cell = table.cell(1, 6)
  cell.text = min_efficiency
  cell = table.cell(1, 7)
  cell.text = min_waso
  cell = table.cell(1, 8)
  cell.text = min_awakenings

  #Row 2 data
  cell = table.cell(2, 1)
  cell.text = max_bedtime
  blank_line = cell.add_paragraph('')
  cell = table.cell(2, 2)
  cell.text = max_get_up_time
  cell = table.cell(2, 3)
  cell.text = max_time_in_bed
  cell = table.cell(2, 4)
  cell.text = max_total_sleep_hours
  cell = table.cell(2, 5)
  cell.text = max_onset_latency
  cell = table.cell(2, 6)
  cell.text = max_efficiency
  cell = table.cell(2, 7)
  cell.text = max_waso
  cell = table.cell(2, 8)
  cell.text = max_awakenings

  #Row 3 data
  cell = table.cell(3, 1)
  cell.text = avg_bedtime
  blank_line = cell.add_paragraph('')
  cell = table.cell(3, 2)
  cell.text = avg_get_up_time
  cell = table.cell(3, 3)
  cell.text = avg_time_in_bed
  cell = table.cell(3, 4)
  cell.text = avg_total_sleep_hours
  cell = table.cell(3, 5)
  cell.text = avg_onset_latency
  cell = table.cell(3, 6)
  cell.text = avg_efficiency
  cell = table.cell(3, 7)
  cell.text = avg_waso
  cell = table.cell(3, 8)
  cell.text = avg_awakenings

  #Formating of table
  style = document.styles['Normal']
  font = style.font
  font.name = 'Arial'
  font.size = Pt(10)
  cell.style = document.styles['Normal']

  #next page
  document.add_page_break()

  #Generating the document
  document.save('Sleep Summary.docx')
